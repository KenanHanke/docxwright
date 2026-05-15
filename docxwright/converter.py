from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    number: int = 0
    rows: list[list[str]] = field(default_factory=list)
    caption: str = ""


@dataclass
class Metadata:
    title: str = ""
    authors: list[str] = field(default_factory=list)
    date: str = ""


@dataclass
class ParseContext:
    figure_count: int = 0
    table_count: int = 0
    labels: dict[str, str] = field(default_factory=dict)


GREEK = {
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "epsilon": "ε",
    "varepsilon": "ε",
    "zeta": "ζ",
    "eta": "η",
    "theta": "θ",
    "vartheta": "ϑ",
    "iota": "ι",
    "kappa": "κ",
    "lambda": "λ",
    "mu": "μ",
    "nu": "ν",
    "xi": "ξ",
    "pi": "π",
    "rho": "ρ",
    "sigma": "σ",
    "tau": "τ",
    "upsilon": "υ",
    "phi": "φ",
    "varphi": "φ",
    "chi": "χ",
    "psi": "ψ",
    "omega": "ω",
    "Gamma": "Γ",
    "Delta": "Δ",
    "Theta": "Θ",
    "Lambda": "Λ",
    "Xi": "Ξ",
    "Pi": "Π",
    "Sigma": "Σ",
    "Phi": "Φ",
    "Psi": "Ψ",
    "Omega": "Ω",
}

SYMBOLS = {
    "times": "×",
    "cdot": "·",
    "pm": "±",
    "le": "≤",
    "leq": "≤",
    "ge": "≥",
    "geq": "≥",
    "neq": "≠",
    "ne": "≠",
    "approx": "≈",
    "sim": "∼",
    "in": "∈",
    "notin": "∉",
    "exists": "∃",
    "forall": "∀",
    "bigvee": "∨",
    "bigcup": "⋃",
    "cup": "∪",
    "cap": "∩",
    "infty": "∞",
    "sum": "∑",
    "sqrt": "√",
    "ln": "ln",
    "bar": "¯",
    "textcopyright": "©",
    "copyright": "©",
    "textregistered": "®",
    "texttrademark": "™",
    "degree": "°",
    "circ": "°",
    "qquad": "  ",
    "quad": " ",
}

SUPERSCRIPT = str.maketrans("0123456789+-=()n", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ")
SUBSCRIPT_MAP = {
    **dict(zip("0123456789+-=()", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎")),
    **dict(zip("aehijklmnoprstuvxy", "ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓᵧ")),
}


def convert_file(tex_path: Path, docx_path: Path, markdown_path: Path | None = None) -> None:
    tex_path = tex_path.resolve()
    docx_path = docx_path.resolve()
    source = flatten_tex(tex_path)
    metadata = extract_metadata(source)
    body = extract_document_body(source)
    context = ParseContext(labels=collect_label_map(body))
    blocks = parse_blocks(body, tex_path.parent, context)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    write_docx(blocks, metadata, docx_path)
    if markdown_path:
        markdown_path = markdown_path.resolve()
        markdown_path.parent.mkdir(parents=True, exist_ok=True)
        markdown_path.write_text(render_markdown(blocks, metadata), encoding="utf-8")


def flatten_tex(path: Path, seen: set[Path] | None = None) -> str:
    seen = seen or set()
    path = path.resolve()
    if path in seen:
        return ""
    seen.add(path)
    text = path.read_text(encoding="utf-8")

    def replace_include(match: re.Match[str]) -> str:
        command, target = match.group(1), match.group(2)
        candidate = (path.parent / target)
        if candidate.suffix != ".tex":
            candidate = candidate.with_suffix(".tex")
        if command == "input" and not candidate.exists():
            return f"\n[Figure content omitted: {target}]\n"
        if candidate.exists():
            return flatten_tex(candidate, seen)
        return ""

    return re.sub(r"\\(subfile|input)\s*\{([^{}]+)\}", replace_include, text)


def extract_metadata(source: str) -> Metadata:
    title = normalize_text(extract_braced_command(source, "title") or "")
    author_text = extract_braced_command(source, "author") or ""
    authors = [normalize_text(part).strip() for part in re.split(r"\\and", author_text) if part.strip()]
    date = extract_braced_command(source, "date") or ""
    if r"\today" in date:
        date = ""
    return Metadata(title=title, authors=authors, date=normalize_text(date))


def extract_document_body(source: str) -> str:
    match = re.search(r"\\begin\{document\}(.*)\\end\{document\}", source, flags=re.S)
    return match.group(1) if match else source


def parse_blocks(text: str, base_dir: Path, context: ParseContext) -> list[Block]:
    text = strip_comments(text)
    text = remove_bibliography(text)
    text = remove_wrappers(text)
    text = re.sub(r"\\maketitle|\\tableofcontents", "\n", text)
    blocks: list[Block] = []
    pos = 0
    env_pattern = re.compile(r"\\begin\{(abstract|figure|table|equation|align|cases)\}(?:\[[^]]*\])?", re.S)
    for match in env_pattern.finditer(text):
        if match.start() < pos:
            continue
        blocks.extend(parse_plain(text[pos : match.start()], context))
        env = match.group(1)
        content_start = match.end()
        end_match = find_environment_end(text, env, content_start)
        if not end_match:
            pos = content_start
            continue
        content = text[content_start : end_match.start()]
        if env == "abstract":
            blocks.append(Block("heading", "Abstract", level=1))
            blocks.extend(parse_plain(content, context))
        elif env == "figure":
            blocks.append(parse_figure(content, context))
        elif env == "table":
            blocks.append(parse_table(content, context))
        else:
            blocks.append(Block("equation", normalize_math(content)))
        pos = end_match.end()
    blocks.extend(parse_plain(text[pos:], context))
    return [block for block in blocks if block.kind != "paragraph" or block.text.strip()]


def parse_plain(text: str, context: ParseContext) -> list[Block]:
    text = re.sub(r"\\end\{document\}|\\begin\{document\}|\\documentclass(?:\[[^]]*\])?\{[^{}]*\}", "\n", text)
    blocks: list[Block] = []
    parts = re.split(r"(\\(?:section|subsection|subsubsection|paragraph)\*?\{[^{}]*\})", text)
    current: list[str] = []
    for part in parts:
        command_match = re.fullmatch(r"\\(section|subsection|subsubsection|paragraph)\*?\{([^{}]*)\}", part, flags=re.S)
        if command_match:
            blocks.extend(paragraph_blocks("".join(current), context))
            current = []
            level = {"section": 1, "subsection": 2, "subsubsection": 3, "paragraph": 4}[command_match.group(1)]
            blocks.append(Block("heading", normalize_text(command_match.group(2), context), level=level))
        else:
            current.append(part)
    blocks.extend(paragraph_blocks("".join(current), context))
    return blocks


def paragraph_blocks(text: str, context: ParseContext) -> list[Block]:
    text = text.replace("\\\\", "\n")
    raw_paragraphs = re.split(r"\n\s*\n+", text)
    blocks = []
    for para in raw_paragraphs:
        para = normalize_text(para, context).strip()
        if para:
            blocks.append(Block("paragraph", para))
    return blocks


def parse_figure(content: str, context: ParseContext) -> Block:
    context.figure_count += 1
    label = extract_braced_command(content, "label")
    if label:
        context.labels[label] = str(context.figure_count)
    caption = normalize_text(extract_braced_command(content, "caption") or "", context)
    text = f"[Figure {context.figure_count} omitted]"
    return Block("figure", text=text, number=context.figure_count, caption=caption)


def parse_table(content: str, context: ParseContext) -> Block:
    context.table_count += 1
    label = extract_braced_command(content, "label")
    if label:
        context.labels[label] = str(context.table_count)
    caption = normalize_text(extract_braced_command(content, "caption") or "", context)
    tabular = extract_environment(content, "tabular")
    rows = parse_tabular(tabular or "", context)
    return Block("table", number=context.table_count, caption=caption, rows=rows)


def parse_tabular(content: str, context: ParseContext) -> list[list[str]]:
    if not content:
        return []
    content = re.sub(r"^\s*\{[^{}]*\}", "", content.strip(), count=1, flags=re.S)
    content = re.sub(r"\\(?:hline|toprule|midrule|bottomrule)\b", "\n", content)
    content = re.sub(r"\\multirow(?:\[[^]]*\])?\{[^{}]*\}\{[^{}]*\}\{([^{}]*)\}", r"\1", content)
    content = re.sub(r"\\multicolumn\{[^{}]*\}\{[^{}]*\}\{([^{}]*)\}", r"\1", content)
    rows = []
    for row_text in split_latex_rows(content):
        cells = [normalize_text(cell, context).strip() for cell in split_latex_cells(row_text)]
        cells = [cell for cell in cells if cell or len(cells) > 1]
        if cells:
            rows.append(cells)
    return rows


def split_latex_rows(text: str) -> list[str]:
    rows: list[str] = []
    buf: list[str] = []
    depth = 0
    i = 0
    while i < len(text):
        ch = text[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth = max(0, depth - 1)
        if depth == 0 and text.startswith(r"\\", i):
            row = "".join(buf).strip()
            if row:
                rows.append(row)
            buf = []
            i += 2
            continue
        buf.append(ch)
        i += 1
    tail = "".join(buf).strip()
    if tail:
        rows.append(tail)
    return rows


def split_latex_cells(row: str) -> list[str]:
    cells: list[str] = []
    buf: list[str] = []
    depth = 0
    for ch in row:
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth = max(0, depth - 1)
        if ch == "&" and depth == 0:
            cells.append("".join(buf))
            buf = []
        else:
            buf.append(ch)
    cells.append("".join(buf))
    return cells


def normalize_text(text: str, context: ParseContext | None = None) -> str:
    context = context or ParseContext()
    text = strip_comments(text)
    text = text.replace("\u2010", "-")
    text = re.sub(r"~|\\[,;:! ]", " ", text)
    text = re.sub(r"\\(?:noindent|centering|small|normalsize|footnotesize|scriptsize)\b", "", text)
    text = text.replace(r"\-", "")
    text = re.sub(r"\\(?:cprotect|protect)\b", "", text)
    text = replace_citations(text)
    text = re.sub(r"\\ref\{([^{}]+)\}", lambda m: context.labels.get(m.group(1), f"reference {m.group(1)}"), text)
    text = re.sub(r"\\label\{[^{}]*\}", "", text)
    text = re.sub(r"\\href\{([^{}]*)\}\{([^{}]*)\}", lambda m: normalize_text(m.group(2), context), text)
    text = re.sub(r"\\url\{([^{}]*)\}|\\nolinkurl\{([^{}]*)\}", lambda m: m.group(1) or m.group(2), text)
    text = replace_math(text)
    text = apply_inline_formatting(text, context)
    text = unwrap_text_commands(text, context)
    for name, value in {**GREEK, **SYMBOLS}.items():
        text = text.replace(f"\\{name}", value)
    text = text.replace("\\%", "%").replace("\\&", "&").replace("\\$", "$").replace("\\#", "#")
    text = text.replace("\\_", "_").replace("\\{", "{").replace("\\}", "}")
    text = text.replace("---", "—").replace("--", "–")
    text = text.replace("``", "“").replace("''", "”")
    text = re.sub(r"\\[a-zA-Z@]+\*?(?:\[[^]]*\])?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    return text.strip()


def apply_inline_formatting(text: str, context: ParseContext) -> str:
    replacements = [
        (r"\\textbf\*?\{([^{}]*)\}", "**{}**"),
        (r"\\(?:textit|emph|mkbibemph)\*?\{([^{}]*)\}", "*{}*"),
        (r"\\(?:texttt|verb)\*?(?:\|([^|]*)\||\{([^{}]*)\})", "`{}`"),
    ]
    previous = None
    while previous != text:
        previous = text
        for pattern, template in replacements:
            text = re.sub(
                pattern,
                lambda m, template=template: template.format(
                    normalize_text((m.group(1) or (m.group(2) if len(m.groups()) > 1 else "")), context)
                ),
                text,
            )
    return text


def replace_citations(text: str) -> str:
    cite_pattern = re.compile(r"\\(?:parencite|cite|textcite|autocite)(?:\[[^]]*\]){0,2}\{([^{}]+)\}")
    return cite_pattern.sub(lambda m: format_citation(m.group(1)), text)


def format_citation(keys: str) -> str:
    citations = []
    for key in keys.split(","):
        key = key.strip()
        year_match = re.search(r"((?:19|20)\d{2})", key)
        year = year_match.group(1) if year_match else "n.d."
        prefix = key[: year_match.start()] if year_match else key
        surname_bits = re.match(r"[a-z\-]+", prefix)
        surname = surname_bits.group(0) if surname_bits else prefix
        surname = "-".join(bit.capitalize() for bit in surname.split("-") if bit)
        if surname == "Muller":
            surname = "Müller"
        citations.append(f"{surname} {year}")
    return "(" + "; ".join(citations) + ")"


def replace_math(text: str) -> str:
    text = re.sub(r"\$\$(.*?)\$\$", lambda m: normalize_math(m.group(1)), text, flags=re.S)
    text = re.sub(r"\$(.*?)\$", lambda m: normalize_math(m.group(1)), text, flags=re.S)
    text = re.sub(r"\\\[(.*?)\\\]", lambda m: normalize_math(m.group(1)), text, flags=re.S)
    text = re.sub(r"\\\((.*?)\\\)", lambda m: normalize_math(m.group(1)), text, flags=re.S)
    return text


def normalize_math(math: str) -> str:
    math = strip_comments(math).strip()
    math = replace_latex_fractions(math)
    math = re.sub(r"\\mathrm\{([^{}]*)\}", r"\1", math)
    math = re.sub(r"\\mathbb\{([^{}]*)\}", r"\1", math)
    math = re.sub(r"\\mathcal\{([^{}]*)\}", r"\1", math)
    math = re.sub(r"\\text\{([^{}]*)\}", r"\1", math)
    math = re.sub(r"\\boldsymbol\{([^{}]*)\}", r"\1", math)
    math = re.sub(r"\\bar\{([^{}])\}", lambda m: m.group(1) + "\u0304", math)
    math = replace_math_commands(math)
    math = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", lambda m: f"{normalize_math(m.group(1))}/{normalize_math(m.group(2))}", math)
    math = re.sub(r"\\sqrt\{([^{}]+)\}", lambda m: f"√({normalize_math(m.group(1))})", math)
    math = re.sub(r"\^\{([^{}]+)\}", lambda m: to_superscript(normalize_math(m.group(1))), math)
    math = re.sub(r"_\{([^{}]+)\}", lambda m: to_subscript(normalize_math(m.group(1))), math)
    math = re.sub(r"\^\\circ\b", "°", math)
    math = re.sub(r"\^([A-Za-z0-9+\-=()])", lambda m: to_superscript(m.group(1)), math)
    math = re.sub(r"_([A-Za-z0-9+\-=()])", lambda m: to_subscript(m.group(1)), math)
    math = re.sub(r"\\begin\{cases\}(.*?)\\end\{cases\}", lambda m: normalize_cases(m.group(1)), math, flags=re.S)
    math = re.sub(r"\\begin\{bmatrix\}(.*?)\\end\{bmatrix\}", lambda m: "[" + normalize_math(m.group(1)) + "]", math, flags=re.S)
    math = math.replace(r"\left", "").replace(r"\right", "")
    math = math.replace(r"\!", "")
    math = re.sub(r"\\(?:,|;|:| )", " ", math)
    math = replace_math_commands(math)
    math = re.sub(r"\\binom\{([^{}]+)\}\{([^{}]+)\}", r"C(\1,\2)", math)
    math = math.replace(r"\mbox{-}", "-")
    math = math.replace(r"\prime", "′")
    math = math.replace(r"\{", "{").replace(r"\}", "}")
    math = math.replace("\\", "")
    math = math.replace("{", "").replace("}", "")
    math = math.replace("---", "—").replace("--", "–")
    math = re.sub(r"\s*&\s*", " ", math)
    math = re.sub(r"\s+", " ", math)
    return math.strip()


def replace_math_commands(math: str) -> str:
    replacements = {**GREEK, **SYMBOLS}
    replacements.update({"vee": "∨", "dots": "…", "ldots": "…", "vdots": "⋮", "mapsto": "↦", "mid": "|"})
    for name, value in sorted(replacements.items(), key=lambda item: -len(item[0])):
        math = re.sub(rf"\\{re.escape(name)}(?=[^A-Za-z]|$)", value, math)
    return math


def normalize_cases(text: str) -> str:
    rows = []
    for row in split_latex_rows(text):
        cells = [normalize_math(cell) for cell in split_latex_cells(row)]
        rows.append(" if ".join(cell for cell in cells if cell))
    return "{ " + "; ".join(rows) + " }"


def replace_latex_fractions(text: str) -> str:
    command = r"\frac"
    start = text.find(command)
    while start != -1:
        first_start = start + len(command)
        first = read_braced_at(text, first_start)
        if not first:
            start = text.find(command, start + len(command))
            continue
        numerator, first_end = first
        second = read_braced_at(text, first_end)
        if not second:
            start = text.find(command, start + len(command))
            continue
        denominator, second_end = second
        replacement = f"({normalize_math(numerator)})/({normalize_math(denominator)})"
        text = text[:start] + replacement + text[second_end:]
        start = text.find(command, start + len(replacement))
    return text


def read_braced_at(text: str, index: int) -> tuple[str, int] | None:
    while index < len(text) and text[index].isspace():
        index += 1
    if index >= len(text) or text[index] != "{":
        return None
    depth = 0
    for pos in range(index, len(text)):
        if text[pos] == "{" and (pos == 0 or text[pos - 1] != "\\"):
            depth += 1
        elif text[pos] == "}" and (pos == 0 or text[pos - 1] != "\\"):
            depth -= 1
            if depth == 0:
                return text[index + 1 : pos], pos + 1
    return None


def to_superscript(value: str) -> str:
    return value.translate(SUPERSCRIPT)


def to_subscript(value: str) -> str:
    return "".join(SUBSCRIPT_MAP.get(char, char) for char in value)


def unwrap_text_commands(text: str, context: ParseContext) -> str:
    command_pattern = re.compile(
        r"\\(?:mathbf|mathit|mathbb|operatorname|underline)\*?(?:\|([^|]*)\||\{([^{}]*)\})"
    )
    previous = None
    while previous != text:
        previous = text
        text = command_pattern.sub(lambda m: normalize_text(m.group(1) or m.group(2) or "", context), text)
    return text


def strip_comments(text: str) -> str:
    return re.sub(r"(?<!\\)%.*", "", text)


def remove_bibliography(text: str) -> str:
    text = re.sub(r"\\printbibliography\b(?:\[[^]]*\])?", "", text)
    text = re.sub(r"\\addbibresource\{[^{}]*\}", "", text)
    text = re.sub(r"\\begin\{thebibliography\}.*?\\end\{thebibliography\}", "", text, flags=re.S)
    return text


def remove_wrappers(text: str) -> str:
    text = re.sub(r"\\begingroup|\\endgroup", "\n", text)
    text = re.sub(r"\\(?:emergencystretch|hfuzz)\s*=\s*[^ \n]+", "", text)
    return text


def collect_label_map(text: str) -> dict[str, str]:
    labels: dict[str, str] = {}
    counts = {"figure": 0, "table": 0}
    pattern = re.compile(r"\\begin\{(figure|table)\}(?:\[[^]]*\])?(.*?)\\end\{\1\}", flags=re.S)
    for match in pattern.finditer(text):
        kind = match.group(1)
        counts[kind] += 1
        label = extract_braced_command(match.group(2), "label")
        if label:
            labels[label] = str(counts[kind])
    return labels


def find_environment_end(text: str, env: str, start: int) -> re.Match[str] | None:
    return re.compile(rf"\\end\{{{re.escape(env)}\}}", flags=re.S).search(text, start)


def extract_environment(text: str, env: str) -> str | None:
    match = re.search(rf"\\begin\{{{re.escape(env)}\}}(.*?)\\end\{{{re.escape(env)}\}}", text, flags=re.S)
    return match.group(1) if match else None


def extract_braced_command(text: str, command: str) -> str | None:
    start = re.search(rf"\\{re.escape(command)}(?:\[[^]]*\])?\{{", text)
    if not start:
        return None
    brace_start = start.end() - 1
    depth = 0
    for idx in range(brace_start, len(text)):
        if text[idx] == "{" and (idx == 0 or text[idx - 1] != "\\"):
            depth += 1
        elif text[idx] == "}" and (idx == 0 or text[idx - 1] != "\\"):
            depth -= 1
            if depth == 0:
                return text[brace_start + 1 : idx]
    return None


def write_docx(blocks: Iterable[Block], metadata: Metadata, path: Path) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    if metadata.title:
        paragraph = document.add_heading(metadata.title, level=0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if metadata.authors:
        paragraph = document.add_paragraph(", ".join(metadata.authors))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if metadata.date:
        paragraph = document.add_paragraph(metadata.date)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for block in blocks:
        if block.kind == "heading":
            document.add_heading(block.text, level=min(max(block.level, 1), 4))
        elif block.kind == "paragraph":
            add_formatted_paragraph(document, block.text)
        elif block.kind == "equation":
            paragraph = document.add_paragraph(block.text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block.kind == "figure":
            paragraph = document.add_paragraph(block.text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if block.caption:
                add_formatted_paragraph(document, f"Figure {block.number}: {block.caption}")
        elif block.kind == "table":
            add_docx_table(document, block)
    document.save(path)


def add_docx_table(document: Document, block: Block) -> None:
    rows = block.rows
    if not rows:
        if block.caption:
            add_formatted_paragraph(document, f"Table {block.number}: {block.caption}")
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""
    if block.caption:
        add_formatted_paragraph(document, f"Table {block.number}: {block.caption}")


def add_formatted_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    add_formatted_runs(paragraph, text)
    return paragraph


def add_formatted_runs(paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in token_pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def render_markdown(blocks: Iterable[Block], metadata: Metadata) -> str:
    lines: list[str] = []
    if metadata.title:
        lines.extend([f"# {metadata.title}", ""])
    if metadata.authors:
        lines.extend([", ".join(metadata.authors), ""])
    for block in blocks:
        if block.kind == "heading":
            lines.extend([f"{'#' * min(block.level + 1, 6)} {block.text}", ""])
        elif block.kind == "paragraph":
            lines.extend([block.text, ""])
        elif block.kind == "equation":
            lines.extend([f"```text\n{block.text}\n```", ""])
        elif block.kind == "figure":
            lines.extend([block.text, ""])
            if block.caption:
                lines.extend([f"Figure {block.number}: {block.caption}", ""])
        elif block.kind == "table":
            lines.extend(render_markdown_table(block.rows))
            if block.caption:
                lines.extend(["", f"Table {block.number}: {block.caption}", ""])
    return "\n".join(lines).rstrip() + "\n"


def render_markdown_table(rows: list[list[str]]) -> list[str]:
    if not rows:
        return []
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(padded[0]) + " |"]
    lines.append("| " + " | ".join("---" for _ in range(width)) + " |")
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines
